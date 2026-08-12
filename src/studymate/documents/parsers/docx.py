"""Word document extraction using python-docx.

DOCX has no native page boundaries, so the whole document is treated
as a single logical page (page_number=1) - splitting further would be
guesswork the format doesn't support.
"""
from __future__ import annotations

from pathlib import Path

import docx

from studymate.documents.parsers.base import ParsedPage


class DOCXParser:
    def parse(self, filepath: Path) -> list[ParsedPage]:
        document = docx.Document(str(filepath))
        parts: list[str] = []
        for para in document.paragraphs:
            if para.text.strip():
                parts.append(para.text)
        for table in document.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        text = "\n".join(parts)
        return [ParsedPage(page_number=1, text=text)]
